"""Case-rapport-generator — DOCX + PDF eksport af én sag.

Producerer en struktureret rapport der samler ALT om en sag:
- Cover-side (sagsnummer, titel, status, dato)
- Executive summary (verdict, krav-antal, evidens-progress)
- Indkøbsproces (intake state)
- Vurderingsmotor-resultat (per-rule decisions, krav, lovcitater)
- Evidens-checkliste (status pr. artefakt + udfyldt indhold)
- Audit-trail (transitions, evidens-completions)
- Lovgrundlag-appendix (alle citerede love)

Bruges af:
- Sag-detalje-side: "Eksportér rapport"-knap → vælg DOCX eller PDF
- V3VurderingPage result-mode: "Download rapport"-knap

Output:
- DOCX (python-docx) — sagsbehandler kan redigere videre
- PDF (reportlab) — print-klar, signerbar version

Sigte: én funktion bygger en data-model (build_report_data); to renderers
(render_docx, render_pdf) producerer outputs fra samme data-model. Sikrer
at de to formater er konsistente.
"""

from __future__ import annotations

import io
import json
import logging
from dataclasses import dataclass, field
from datetime import datetime, UTC
from typing import Any, Optional

from sqlalchemy.orm import Session

from src.database.cases import Case, CaseTransition, find_case_by_external_id
from src.database.evidence import list_evidence_for_case
from src.rule_engine import audit as v3_audit
from src.services.evidence_artifacts import get_template

logger = logging.getLogger(__name__)


# ---- Data model -----------------------------------------------------------


@dataclass
class ReportRuleDecision:
    rule_id: str
    status: str
    lov: str
    artikel: str
    citat: str
    url: str
    begrundelse: str
    krav: list[str] = field(default_factory=list)
    evidens_paakraevet: list[str] = field(default_factory=list)


@dataclass
class ReportEvidence:
    artifact_id: str
    title: str
    status: str  # mangler/i_gang/faerdig/godkendt
    completed_at: Optional[str]
    completed_by: Optional[str]
    sections: list[dict] = field(default_factory=list)  # {key, heading, value, required}


@dataclass
class ReportEvent:
    timestamp: str
    label: str
    detail: str
    actor: Optional[str]
    kind: str


@dataclass
class CaseReportData:
    """Komplet rapport-data — én sag samlet."""

    # Header
    case_id: str
    title: str
    status: str
    status_label: str
    last_aggregate_status: Optional[str]
    assigned_to: Optional[str]
    created_at: Optional[str]
    updated_at: Optional[str]
    next_review_at: Optional[str]
    generated_at: str

    # Intake (indkøbsproces)
    intake_state: dict

    # Latest assessment
    latest_assessment_log_id: Optional[str]
    latest_assessment_at: Optional[str]
    rule_engine_version: Optional[str]
    decisions: list[ReportRuleDecision] = field(default_factory=list)
    total_krav: int = 0
    total_artefakter: int = 0

    # Evidence
    evidence: list[ReportEvidence] = field(default_factory=list)
    evidence_done: int = 0
    evidence_total: int = 0

    # Audit-trail
    events: list[ReportEvent] = field(default_factory=list)


# ---- Build data from DB ---------------------------------------------------


def build_report_data(session: Session, case_id: str) -> Optional[CaseReportData]:
    """Hent + samme alle data om sagen til en struktureret model."""
    case = find_case_by_external_id(session, case_id)
    if case is None:
        return None

    # Latest assessment for this case (filtered by external case_id)
    latest_assessment = None
    decisions: list[ReportRuleDecision] = []
    total_krav = 0
    total_artefakter = set()

    try:
        recent = v3_audit.list_recent(session, limit=10, case_id=case_id)
        if recent:
            latest_assessment = recent[0]
            try:
                response_payload = latest_assessment.response_payload
                if isinstance(response_payload, str):
                    response_payload = json.loads(response_payload)
                # Iterate decisions in the assessment's response
                for d in response_payload.get("decisions", []):
                    if not d.get("triggered"):
                        continue
                    status = d.get("status")
                    if status not in ("BETINGET-GO", "NO-GO"):
                        continue
                    outcome = d.get("outcome") or {}
                    kilde = d.get("kilde") or {}
                    krav = outcome.get("krav") or []
                    artefakter = outcome.get("evidens_påkrævet") or []
                    total_krav += len(krav)
                    total_artefakter.update(artefakter)
                    decisions.append(ReportRuleDecision(
                        rule_id=d.get("rule_id", ""),
                        status=status,
                        lov=kilde.get("lov", ""),
                        artikel=kilde.get("artikel", ""),
                        citat=kilde.get("citat", ""),
                        url=kilde.get("url", ""),
                        begrundelse=outcome.get("begrundelse", ""),
                        krav=krav,
                        evidens_paakraevet=artefakter,
                    ))
            except Exception as exc:
                logger.warning(f"failed to parse assessment payload: {exc}")
    except Exception as exc:
        logger.warning(f"failed to fetch assessment for {case_id}: {exc}")

    # Evidence rows
    evidence_rows = []
    evidence_done = 0
    try:
        rows = list_evidence_for_case(session, case_id)
        for row in rows:
            tmpl = get_template(row.artifact_id)
            content = row.get_content()
            sections_data = []
            for section in tmpl.sections:
                value = content.get(section.key)
                sections_data.append({
                    "key": section.key,
                    "heading": section.heading,
                    "value": value if value not in (None, "") else None,
                    "required": section.required,
                })
            evidence_rows.append(ReportEvidence(
                artifact_id=row.artifact_id,
                title=tmpl.title,
                status=row.status,
                completed_at=row.completed_at.isoformat() if row.completed_at else None,
                completed_by=row.completed_by,
                sections=sections_data,
            ))
            if row.status in ("faerdig", "godkendt"):
                evidence_done += 1
    except Exception as exc:
        logger.warning(f"failed to fetch evidence for {case_id}: {exc}")

    # Audit-trail (transitions)
    events: list[ReportEvent] = []
    try:
        transitions = (
            session.query(CaseTransition)
            .filter(CaseTransition.case_db_id == case.id)
            .order_by(CaseTransition.changed_at.desc())
            .all()
        )
        for t in transitions:
            events.append(ReportEvent(
                timestamp=t.changed_at.isoformat() if t.changed_at else "",
                label=f"Status: {t.from_status or '(ny)'} → {t.to_status}",
                detail=t.note or "",
                actor=t.changed_by,
                kind="transition",
            ))
    except Exception as exc:
        logger.warning(f"failed to fetch transitions: {exc}")

    # Add evidence events
    for ev in evidence_rows:
        if ev.completed_at:
            events.append(ReportEvent(
                timestamp=ev.completed_at,
                label=f"Evidens færdig: {ev.title}",
                detail="",
                actor=ev.completed_by,
                kind="evidens",
            ))

    events.sort(key=lambda e: e.timestamp, reverse=True)

    intake = case.get_intake_state() or {}
    from src.database.cases import CASE_STATUS_LABELS

    return CaseReportData(
        case_id=case.case_id,
        title=case.title or case.case_id,
        status=case.status,
        status_label=CASE_STATUS_LABELS.get(case.status, case.status),
        last_aggregate_status=case.last_aggregate_status,
        assigned_to=case.assigned_to,
        created_at=case.created_at.isoformat() if case.created_at else None,
        updated_at=case.updated_at.isoformat() if case.updated_at else None,
        next_review_at=case.next_review_at.isoformat() if case.next_review_at else None,
        generated_at=datetime.now(UTC).isoformat(),
        intake_state=intake,
        latest_assessment_log_id=latest_assessment.id if latest_assessment else None,
        latest_assessment_at=latest_assessment.created_at.isoformat() if latest_assessment and latest_assessment.created_at else None,
        rule_engine_version=getattr(latest_assessment, "rule_engine_version", None) if latest_assessment else None,
        decisions=decisions,
        total_krav=total_krav,
        total_artefakter=len(total_artefakter),
        evidence=evidence_rows,
        evidence_done=evidence_done,
        evidence_total=len(evidence_rows),
        events=events,
    )


# ---- Helpers --------------------------------------------------------------


def _format_date(iso: Optional[str]) -> str:
    if not iso:
        return "—"
    try:
        d = datetime.fromisoformat(iso.replace("Z", "+00:00"))
        return d.strftime("%-d. %b %Y kl. %H:%M")
    except Exception:
        return iso


def _format_indkoeb_value(key: str, value: Any) -> str:
    """Lav læsbar tekst af intake_state-værdier."""
    if value is None or value == "":
        return "—"
    if isinstance(value, bool):
        return "Ja" if value else "Nej"
    if isinstance(value, str):
        return value.replace("_", " ")
    return str(value)


_INTAKE_FIELD_ORDER = [
    ("behov", "Behovsbeskrivelse"),
    ("dobbeltsystem_tjekket", "Dobbeltsystem-tjekket"),
    ("sagsnummer", "Sagsnummer (Serviceportal)"),
    ("serviceportal_dato", "Oprettelsesdato"),
    ("indkoeb_eller_udvikling", "Indkøb eller udvikling"),
    ("system_description", "Systembeskrivelse"),
    ("current_step", "Aktuelt trin"),
]


# ---- DOCX renderer --------------------------------------------------------


def render_docx(data: CaseReportData) -> bytes:
    """Producér rapport som DOCX. Sagsbehandler kan redigere videre."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ---- Page setup ----
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ---- Cover ----
    eyebrow = doc.add_paragraph()
    eyebrow.add_run("BIFROST · KOMMUNAL AI-COMPLIANCE-RAPPORT").font.size = Pt(8)
    eyebrow.runs[0].font.color.rgb = RGBColor(0xB0, 0x8A, 0x4A)  # bronze
    eyebrow.runs[0].bold = True

    title = doc.add_paragraph()
    title_run = title.add_run(data.title)
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(0x14, 0x18, 0x1F)
    title_run.bold = True

    meta = doc.add_paragraph()
    meta_text = f"Sag {data.case_id}"
    if data.assigned_to:
        meta_text += f" · {data.assigned_to}"
    meta_text += f" · {data.status_label}"
    meta.add_run(meta_text).font.size = Pt(11)

    if data.last_aggregate_status:
        verdict = doc.add_paragraph()
        verdict_run = verdict.add_run(f"Verdict: {data.last_aggregate_status}")
        verdict_run.font.size = Pt(14)
        verdict_run.bold = True
        color_map = {
            "GO": RGBColor(0x2D, 0x6A, 0x31),
            "BETINGET-GO": RGBColor(0xB0, 0x8A, 0x4A),
            "NO-GO": RGBColor(0xA0, 0x20, 0x20),
        }
        if data.last_aggregate_status in color_map:
            verdict_run.font.color.rgb = color_map[data.last_aggregate_status]

    doc.add_paragraph(f"Genereret {_format_date(data.generated_at)}").runs[0].font.size = Pt(9)

    doc.add_paragraph()  # spacer

    # ---- Executive summary ----
    _add_h2(doc, "Sammendrag")
    summary_table = doc.add_table(rows=4, cols=2)
    summary_table.style = "Light List Accent 1"
    summary_table.cell(0, 0).text = "Status"
    summary_table.cell(0, 1).text = data.status_label
    summary_table.cell(1, 0).text = "Verdict"
    summary_table.cell(1, 1).text = data.last_aggregate_status or "—"
    summary_table.cell(2, 0).text = "Evidens-progress"
    summary_table.cell(2, 1).text = f"{data.evidence_done}/{data.evidence_total}"
    summary_table.cell(3, 0).text = "Antal krav"
    summary_table.cell(3, 1).text = str(data.total_krav)

    # ---- Indkøbsproces ----
    if data.intake_state:
        _add_h2(doc, "1. Indkøbsproces")
        intake_table = doc.add_table(rows=0, cols=2)
        intake_table.style = "Light List Accent 1"
        for key, label in _INTAKE_FIELD_ORDER:
            value = data.intake_state.get(key)
            if value is None or value == "":
                continue
            row = intake_table.add_row()
            row.cells[0].text = label
            row.cells[1].text = _format_indkoeb_value(key, value)

    # ---- Vurdering ----
    _add_h2(doc, "2. Bifrost-vurdering")
    if data.latest_assessment_at:
        doc.add_paragraph(
            f"Vurderet: {_format_date(data.latest_assessment_at)} · "
            f"rule_engine {data.rule_engine_version or '?'}"
        ).runs[0].font.size = Pt(9)

    if not data.decisions:
        doc.add_paragraph("Ingen vurdering kørt på sagen endnu.").italic = True
    else:
        for i, d in enumerate(data.decisions, 1):
            _add_h3(doc, f"2.{i} {d.lov} — {d.artikel}")

            status_p = doc.add_paragraph()
            status_run = status_p.add_run(d.status)
            status_run.bold = True
            color_map = {
                "BETINGET-GO": RGBColor(0xB0, 0x8A, 0x4A),
                "NO-GO": RGBColor(0xA0, 0x20, 0x20),
            }
            if d.status in color_map:
                status_run.font.color.rgb = color_map[d.status]

            if d.begrundelse:
                doc.add_paragraph(d.begrundelse)

            if d.citat:
                citat_p = doc.add_paragraph()
                citat_run = citat_p.add_run(f'"{d.citat}"')
                citat_run.italic = True
                citat_run.font.size = Pt(10)

            if d.krav:
                _add_p_label(doc, "Krav:")
                for k in d.krav:
                    bullet = doc.add_paragraph(k, style="List Bullet")
                    bullet.runs[0].font.size = Pt(10.5)

            if d.evidens_paakraevet:
                _add_p_label(doc, "Påkrævet evidens:")
                for e in d.evidens_paakraevet:
                    bullet = doc.add_paragraph(e.replace("_", " "), style="List Bullet")
                    bullet.runs[0].font.size = Pt(10.5)

            if d.url:
                doc.add_paragraph(f"Kilde: {d.url}").runs[0].font.size = Pt(8)

    # ---- Evidens ----
    _add_h2(doc, "3. Evidens-checkliste")
    if not data.evidence:
        doc.add_paragraph("Ingen evidens-artefakter registreret endnu.").italic = True
    else:
        for ev in data.evidence:
            status_label = {
                "mangler": "Mangler",
                "i_gang": "I gang",
                "faerdig": "Færdig",
                "godkendt": "Godkendt",
            }.get(ev.status, ev.status)
            _add_h3(doc, f"{ev.title} — {status_label}")

            if ev.completed_at:
                doc.add_paragraph(
                    f"Færdig {_format_date(ev.completed_at)}{f' af {ev.completed_by}' if ev.completed_by else ''}"
                ).runs[0].font.size = Pt(8)

            for s in ev.sections:
                if s["value"] in (None, ""):
                    continue
                _add_p_label(doc, s["heading"])
                value = s["value"]
                if isinstance(value, bool):
                    value = "Ja" if value else "Nej"
                doc.add_paragraph(str(value)).runs[0].font.size = Pt(10.5)

    # ---- Audit-trail ----
    if data.events:
        _add_h2(doc, "4. Audit-trail")
        events_table = doc.add_table(rows=1, cols=3)
        events_table.style = "Light List Accent 1"
        hdr = events_table.rows[0]
        hdr.cells[0].text = "Tidspunkt"
        hdr.cells[1].text = "Hændelse"
        hdr.cells[2].text = "Aktør"
        for e in data.events[:50]:
            row = events_table.add_row()
            row.cells[0].text = _format_date(e.timestamp)
            row.cells[1].text = e.label
            row.cells[2].text = e.actor or "—"

    # ---- Footer ----
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        "Bifrost — Kalundborg Kommune · Genereret automatisk fra sagens "
        "registrerede data. Sagsbehandler skal verificere indhold før "
        "underskrift eller deling."
    )
    footer_run.font.size = Pt(8)
    footer_run.italic = True
    footer_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # Output to bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_h2(doc, text: str):
    from docx.shared import Pt, RGBColor
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(0x0D, 0x2E, 0x54)


def _add_h3(doc, text: str):
    from docx.shared import Pt, RGBColor
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = RGBColor(0x14, 0x18, 0x1F)


def _add_p_label(doc, text: str):
    from docx.shared import Pt, RGBColor
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.bold = True
    run.font.color.rgb = RGBColor(0x55, 0x5A, 0x64)


# ---- PDF renderer ---------------------------------------------------------


def render_pdf(data: CaseReportData) -> bytes:
    """Producér rapport som PDF via reportlab. Print-klar."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, KeepTogether,
    )
    from reportlab.lib.enums import TA_LEFT

    BRONZE = colors.HexColor("#b08a4a")
    NAVY = colors.HexColor("#0d2e54")
    INK = colors.HexColor("#14181f")
    INK_SOFT = colors.HexColor("#555a64")
    SUCCESS = colors.HexColor("#2d6a31")
    DANGER = colors.HexColor("#a02020")
    LINE = colors.HexColor("#d8d3c5")

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        topMargin=2.2 * cm, bottomMargin=2.2 * cm,
        leftMargin=2.2 * cm, rightMargin=2.2 * cm,
        title=f"Bifrost-rapport — {data.case_id}",
        author=data.assigned_to or "Bifrost",
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="Eyebrow", parent=styles["Normal"], fontSize=8, textColor=BRONZE, spaceAfter=4, leading=10))
    styles.add(ParagraphStyle(name="HeroTitle", parent=styles["Title"], fontSize=24, textColor=INK, spaceAfter=8, leading=28))
    styles.add(ParagraphStyle(name="Meta", parent=styles["Normal"], fontSize=10, textColor=INK_SOFT, spaceAfter=10))
    styles.add(ParagraphStyle(name="VerdictLarge", parent=styles["Normal"], fontSize=14, spaceAfter=12))
    styles.add(ParagraphStyle(name="H2Bifrost", parent=styles["Heading2"], fontSize=14, textColor=NAVY, spaceBefore=18, spaceAfter=8))
    styles.add(ParagraphStyle(name="H3Bifrost", parent=styles["Heading3"], fontSize=11.5, textColor=INK, spaceBefore=10, spaceAfter=4))
    styles.add(ParagraphStyle(name="Citat", parent=styles["Italic"], fontSize=9.5, textColor=INK_SOFT, leftIndent=12, leading=12))
    styles.add(ParagraphStyle(name="LabelSmall", parent=styles["Normal"], fontSize=8, textColor=INK_SOFT, spaceBefore=6, spaceAfter=2))
    styles.add(ParagraphStyle(name="Body", parent=styles["Normal"], fontSize=10, leading=14, spaceAfter=4))
    styles.add(ParagraphStyle(name="Small", parent=styles["Normal"], fontSize=8, textColor=INK_SOFT, leading=10))
    styles.add(ParagraphStyle(name="Footer", parent=styles["Italic"], fontSize=7.5, textColor=colors.grey, leading=9))

    story = []

    # ---- Cover ----
    story.append(Paragraph("BIFROST · KOMMUNAL AI-COMPLIANCE-RAPPORT", styles["Eyebrow"]))
    story.append(Paragraph(_xml_safe(data.title), styles["HeroTitle"]))
    meta_parts = [f"Sag <b>{data.case_id}</b>"]
    if data.assigned_to:
        meta_parts.append(data.assigned_to)
    meta_parts.append(data.status_label)
    story.append(Paragraph(" · ".join(meta_parts), styles["Meta"]))

    if data.last_aggregate_status:
        verdict_color = {
            "GO": SUCCESS,
            "BETINGET-GO": BRONZE,
            "NO-GO": DANGER,
        }.get(data.last_aggregate_status, INK)
        story.append(Paragraph(
            f'<font color="{verdict_color.hexval()}"><b>Verdict: {data.last_aggregate_status}</b></font>',
            styles["VerdictLarge"],
        ))

    story.append(Paragraph(f"Genereret {_format_date(data.generated_at)}", styles["Small"]))
    story.append(Spacer(1, 0.4 * cm))

    # ---- Executive summary ----
    story.append(Paragraph("Sammendrag", styles["H2Bifrost"]))
    summary_data = [
        ["Status", data.status_label],
        ["Verdict", data.last_aggregate_status or "—"],
        ["Evidens-progress", f"{data.evidence_done} / {data.evidence_total}"],
        ["Antal krav", str(data.total_krav)],
    ]
    summary_table = Table(summary_data, colWidths=[5 * cm, None])
    summary_table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("TEXTCOLOR", (0, 0), (0, -1), INK_SOFT),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("LINEBELOW", (0, 0), (-1, -2), 0.3, LINE),
    ]))
    story.append(summary_table)

    # ---- 1. Indkøbsproces ----
    if data.intake_state:
        story.append(Paragraph("1. Indkøbsproces", styles["H2Bifrost"]))
        intake_rows = []
        for key, label in _INTAKE_FIELD_ORDER:
            value = data.intake_state.get(key)
            if value is None or value == "":
                continue
            intake_rows.append([
                Paragraph(label, styles["LabelSmall"]),
                Paragraph(_xml_safe(_format_indkoeb_value(key, value)), styles["Body"]),
            ])
        if intake_rows:
            t = Table(intake_rows, colWidths=[5 * cm, None])
            t.setStyle(TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("LINEBELOW", (0, 0), (-1, -2), 0.2, LINE),
            ]))
            story.append(t)

    # ---- 2. Vurdering ----
    story.append(Paragraph("2. Bifrost-vurdering", styles["H2Bifrost"]))
    if data.latest_assessment_at:
        story.append(Paragraph(
            f"Vurderet: {_format_date(data.latest_assessment_at)} · rule_engine {data.rule_engine_version or '?'}",
            styles["Small"],
        ))

    if not data.decisions:
        story.append(Paragraph("<i>Ingen vurdering kørt på sagen endnu.</i>", styles["Body"]))
    else:
        for i, d in enumerate(data.decisions, 1):
            block = []
            block.append(Paragraph(f"2.{i} {_xml_safe(d.lov)} — {_xml_safe(d.artikel)}", styles["H3Bifrost"]))
            verdict_color = BRONZE if d.status == "BETINGET-GO" else DANGER
            block.append(Paragraph(
                f'<font color="{verdict_color.hexval()}"><b>{d.status}</b></font>',
                styles["Body"],
            ))
            if d.begrundelse:
                block.append(Paragraph(_xml_safe(d.begrundelse), styles["Body"]))
            if d.citat:
                block.append(Paragraph(f'"{_xml_safe(d.citat)}"', styles["Citat"]))
            if d.krav:
                block.append(Paragraph("Krav:", styles["LabelSmall"]))
                for k in d.krav:
                    block.append(Paragraph(f"• {_xml_safe(k)}", styles["Body"]))
            if d.evidens_paakraevet:
                block.append(Paragraph("Påkrævet evidens:", styles["LabelSmall"]))
                for e in d.evidens_paakraevet:
                    block.append(Paragraph(f"• {_xml_safe(e.replace('_', ' '))}", styles["Body"]))
            if d.url:
                block.append(Paragraph(f"Kilde: <link href=\"{d.url}\">{d.url}</link>", styles["Small"]))
            story.append(KeepTogether(block))
            story.append(Spacer(1, 0.2 * cm))

    # ---- 3. Evidens ----
    story.append(Paragraph("3. Evidens-checkliste", styles["H2Bifrost"]))
    if not data.evidence:
        story.append(Paragraph("<i>Ingen evidens-artefakter registreret endnu.</i>", styles["Body"]))
    else:
        for ev in data.evidence:
            status_label = {
                "mangler": "Mangler",
                "i_gang": "I gang",
                "faerdig": "Færdig",
                "godkendt": "Godkendt",
            }.get(ev.status, ev.status)
            block = [Paragraph(
                f"{_xml_safe(ev.title)} — <i>{status_label}</i>",
                styles["H3Bifrost"],
            )]
            if ev.completed_at:
                block.append(Paragraph(
                    f"Færdig {_format_date(ev.completed_at)}"
                    f"{f' af {ev.completed_by}' if ev.completed_by else ''}",
                    styles["Small"],
                ))
            for s in ev.sections:
                if s["value"] in (None, ""):
                    continue
                block.append(Paragraph(_xml_safe(s["heading"]), styles["LabelSmall"]))
                value = s["value"]
                if isinstance(value, bool):
                    value = "Ja" if value else "Nej"
                block.append(Paragraph(_xml_safe(str(value)), styles["Body"]))
            story.append(KeepTogether(block))
            story.append(Spacer(1, 0.15 * cm))

    # ---- 4. Audit-trail ----
    if data.events:
        story.append(Paragraph("4. Audit-trail", styles["H2Bifrost"]))
        rows = [["Tidspunkt", "Hændelse", "Aktør"]]
        for e in data.events[:50]:
            rows.append([
                _format_date(e.timestamp),
                _xml_safe(e.label),
                e.actor or "—",
            ])
        t = Table(rows, colWidths=[4.2 * cm, 9 * cm, 3 * cm])
        t.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8.5),
            ("TEXTCOLOR", (0, 0), (-1, 0), INK_SOFT),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f5f4ef")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("LINEBELOW", (0, 0), (-1, -1), 0.2, LINE),
        ]))
        story.append(t)

    # ---- Footer ----
    story.append(Spacer(1, 0.6 * cm))
    story.append(Paragraph(
        "Bifrost — Kalundborg Kommune · Genereret automatisk fra sagens registrerede data. "
        "Sagsbehandler skal verificere indhold før underskrift eller deling.",
        styles["Footer"],
    ))

    doc.build(story, onFirstPage=_pdf_footer, onLaterPages=_pdf_footer)
    return buf.getvalue()


def _pdf_footer(canvas, doc):
    """Skriv side-nummer + brand i bunden af hver side."""
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    canvas.saveState()
    canvas.setFont("Helvetica", 7.5)
    canvas.setFillColor(colors.HexColor("#8a8f96"))
    canvas.drawString(2.2 * cm, 1.2 * cm, "Bifrost ᛒ · kommunal AI-compliance · Kalundborg Kommune")
    canvas.drawRightString(doc.pagesize[0] - 2.2 * cm, 1.2 * cm, f"Side {canvas.getPageNumber()}")
    canvas.restoreState()


def _xml_safe(text: str) -> str:
    """Escape characters that confuse reportlab's mini-XML parser."""
    if text is None:
        return ""
    return (
        str(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )
