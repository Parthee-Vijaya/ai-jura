"""Tests for document-analyse pipeline (M1)."""

from __future__ import annotations

import io

import pytest

from src.services.document_analyzer import (
    parse_pdf,
    parse_docx,
    parse_document,
    chunk_text,
    _merge_signals,
    _find_page,
    Chunk,
    ChunkSignals,
    DocumentParseError,
)


def _make_docx_bytes(headings_paragraphs: list[tuple[str, str]]) -> bytes:
    from docx import Document
    doc = Document()
    for heading, body in headings_paragraphs:
        if heading:
            doc.add_heading(heading, 1)
        if body:
            doc.add_paragraph(body)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestChunking:
    def test_short_text_one_chunk(self):
        text = "Kort tekst der ikke kræver opsplitning."
        chunks = chunk_text(text, [], chunk_size=4000, overlap=400)
        assert len(chunks) == 1
        assert chunks[0].index == 0
        assert chunks[0].char_start == 0

    def test_long_text_splits_with_overlap(self):
        text = "Dette er en sætning. " * 500  # ~10500 chars
        chunks = chunk_text(text, [], chunk_size=2000, overlap=200)
        assert len(chunks) > 1
        # Each chunk respects size cap (with breaking-point flexibility)
        for c in chunks:
            assert len(c.text) <= 2000
        # Last chunk reaches end of text
        assert chunks[-1].char_end >= len(text) - 100

    def test_chunk_index_is_stable_and_sequential(self):
        text = "x" * 5000
        chunks = chunk_text(text, [], chunk_size=1000, overlap=100)
        for i, c in enumerate(chunks):
            assert c.index == i

    def test_empty_text_returns_no_chunks(self):
        chunks = chunk_text("", [], chunk_size=4000, overlap=400)
        assert chunks == []

    def test_chunks_cover_full_text(self):
        text = "Section one. " * 30 + "\n\n" + "Section two. " * 30
        chunks = chunk_text(text, [], chunk_size=500, overlap=100)
        # Chunks should collectively cover the full document
        coverage = "".join(c.text for c in chunks)
        # Allow for whitespace differences from chunk normalization
        assert "Section one" in coverage
        assert "Section two" in coverage


class TestPageMapping:
    def test_find_page_within_offset(self):
        offsets = [(1, 0, 100), (2, 100, 200), (3, 200, 300)]
        assert _find_page(50, offsets) == 1
        assert _find_page(150, offsets) == 2
        assert _find_page(250, offsets) == 3

    def test_find_page_after_last_uses_last_offset(self):
        offsets = [(1, 0, 100), (2, 100, 200)]
        # 500 is past last offset's end — falls back to last page label
        assert _find_page(500, offsets) == 2

    def test_find_page_no_offsets_returns_none(self):
        assert _find_page(50, []) is None


class TestSignalMerging:
    def test_any_true_wins(self):
        cs1 = ChunkSignals(chunk=Chunk(0, "x"), signals={"a": False})
        cs2 = ChunkSignals(chunk=Chunk(1, "y"), signals={"a": True})
        merged = _merge_signals([cs1, cs2])
        assert merged["a"] is True

    def test_false_persists_when_no_chunk_says_true(self):
        cs1 = ChunkSignals(chunk=Chunk(0, "x"), signals={"b": False})
        cs2 = ChunkSignals(chunk=Chunk(1, "y"), signals={"b": False})
        merged = _merge_signals([cs1, cs2])
        assert merged["b"] is False

    def test_merging_disjoint_signals(self):
        cs1 = ChunkSignals(chunk=Chunk(0, "x"), signals={"a": True})
        cs2 = ChunkSignals(chunk=Chunk(1, "y"), signals={"b": True})
        merged = _merge_signals([cs1, cs2])
        assert merged == {"a": True, "b": True}

    def test_empty_chunk_signals(self):
        cs = ChunkSignals(chunk=Chunk(0, "x"), signals={})
        merged = _merge_signals([cs])
        assert merged == {}


class TestParsing:
    def test_parse_docx_returns_text_and_offsets(self):
        docx_bytes = _make_docx_bytes(
            [("Heading", "First paragraph."), ("", "Second paragraph.")]
        )
        text, offsets = parse_docx(docx_bytes)
        assert "First paragraph" in text
        assert "Second paragraph" in text
        assert len(offsets) >= 2

    def test_parse_document_detects_docx_via_filename(self):
        docx_bytes = _make_docx_bytes([("h", "body")])
        text, offsets, kind = parse_document(docx_bytes, "test.docx")
        assert kind == "docx"
        assert "body" in text

    def test_parse_document_unsupported_format_raises(self):
        with pytest.raises(DocumentParseError, match="Unsupported file type"):
            parse_document(b"plain text content", "noteopen.txt")

    def test_parse_empty_docx(self):
        # Empty docx with no paragraphs — shouldn't crash
        from docx import Document
        doc = Document()
        buf = io.BytesIO()
        doc.save(buf)
        text, offsets = parse_docx(buf.getvalue())
        assert text == ""
        assert offsets == []
